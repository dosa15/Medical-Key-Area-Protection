import docx
from docx.shared import Cm
import zipfile
import random
import sys

import numpy as np
from PIL import Image
from cv2 import medianBlur 

cmppath = sys.argv[1]
destpath = './tmp/'

if __name__ == '__main__':
    image1 = Image.open(destpath + 'compressed/' + cmppath)
    image1_array = np.asarray(image1)

    decryption_key = np.load('key.npy')
    
    docText = ''
    for dni in decryption_key:
        # print(image1_array[dni[0], dni[1]])
        docText += chr(image1_array[dni[0], dni[1], 0])
    print(docText)

    # output = cv2.fastNlMeansDenoisingColored(image1_array, None, 2, 2, 7, 21)
    output = Image.fromarray(medianBlur(image1_array, 3))
    output.show()
    output.save(destpath + "uncompressed/ucmp1_image1.tif")

    uncompressed_docx = docx.Document()
    for para in docText.split('\n\n'):
        uncompressed_docx.add_paragraph(para)
    uncompressed_docx.add_picture(destpath + "uncompressed/ucmp1_image1.tif", width=Cm(16.5))
    uncompressed_docx.save('test_uncompressed.docx')